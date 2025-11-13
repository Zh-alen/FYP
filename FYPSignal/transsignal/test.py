#create test signal
# create_test_files.py
import os
from docx import Document
import openpyxl

# Create test Word document
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test Word document for binary conversion demonstration.')
doc.add_paragraph('Contains some sample text data.')
doc.save('C:/SJTU-simcom/FYPSignal/input_files/test_document.docx')

# Create test Excel file
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Data Sheet"
ws['A1'] = 'Name'
ws['B1'] = 'Age'
ws['A2'] = 'John'
ws['B2'] = 25
ws['A3'] = 'Jane'
ws['B3'] = 30
wb.save('C:/SJTU-simcom/FYPSignal/input_files/test_data.xlsx')