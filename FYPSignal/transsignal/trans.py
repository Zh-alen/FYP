import json
from datetime import datetime, date
import os
import numpy as np
import jax.numpy as jnp
from docx import Document
import openpyxl

class CustomJSONEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, (datetime, date)):
            return obj.isoformat()
        elif hasattr(obj, '__dict__'):
            return obj.__dict__
        return super().default(obj)

class AdvancedBinaryEncoder:
    def __init__(self):
        self.encoding_methods = {
            'direct_binary': self.direct_binary_encode,
            'manchester': self.manchester_encode,
            'nrzi': self.nrzi_encode,
            'hamming_74': self.hamming_74_encode,  # 修复方法名
            'polar_basic': self.polar_basic_encode
        }
        self.json_encoder = CustomJSONEncoder()
    
    def direct_binary_encode(self, data_bytes):
        """直接二进制编码"""
        binary_string = ''.join(format(byte, '08b') for byte in data_bytes)
        return binary_string, "Direct binary encoding"
    
    def manchester_encode(self, data_bytes):
        """曼彻斯特编码"""
        binary_string = ''.join(format(byte, '08b') for byte in data_bytes)
        manchester_encoded = ''.join('01' if bit == '0' else '10' for bit in binary_string)
        return manchester_encoded, "Manchester encoding"
    
    def nrzi_encode(self, data_bytes):
        """NRZI编码"""
        binary_string = ''.join(format(byte, '08b') for byte in data_bytes)
        nrzi_encoded = ''
        current_level = '1'
        for bit in binary_string:
            if bit == '0':
                current_level = '0' if current_level == '1' else '1'
            nrzi_encoded += current_level
        return nrzi_encoded, "NRZI encoding"
    
    def hamming_74_encode(self, data_bytes, chunk_size=1000):
        """汉明码(7,4)编码 - 修复方法名"""
        # 对于小文件直接处理，大文件使用流式处理
        if len(data_bytes) <= 10000:
            binary_string = ''.join(format(byte, '08b') for byte in data_bytes)
            encoded_string = self._hamming_encode_bits(binary_string)
            return encoded_string, "Hamming(7,4) encoding"
        else:
            return self.hamming_74_encode_streaming(data_bytes, chunk_size)
    
    def hamming_74_encode_streaming(self, data_bytes, chunk_size=1000):
        """汉明码(7,4)编码 - 流式处理"""
        total_bits = len(data_bytes) * 8
        encoded_chunks = []
        
        for i in range(0, len(data_bytes), chunk_size):
            chunk = data_bytes[i:i + chunk_size]
            binary_string = ''.join(format(byte, '08b') for byte in chunk)
            encoded_chunk = self._hamming_encode_bits(binary_string)
            encoded_chunks.append(encoded_chunk)
        
        return ''.join(encoded_chunks), "Hamming(7,4) encoding (streaming)"
    
    def _hamming_encode_bits(self, binary_string):
        """对二进制字符串进行汉明编码"""
        hamming_encoded = ''
        
        for i in range(0, len(binary_string), 4):
            data_bits = binary_string[i:i+4]
            if len(data_bits) < 4:
                data_bits = data_bits + '0' * (4 - len(data_bits))
            
            d1, d2, d3, d4 = [int(bit) for bit in data_bits]
            p1 = d1 ^ d2 ^ d4
            p2 = d1 ^ d3 ^ d4
            p3 = d2 ^ d3 ^ d4
            
            encoded_bits = [p1, p2, d1, p3, d2, d3, d4]
            hamming_encoded += ''.join(str(bit) for bit in encoded_bits)
        
        return hamming_encoded
    
    def polar_basic_encode(self, data_bytes, n=8):
        """基础Polar编码"""
        binary_string = ''.join(format(byte, '08b') for byte in data_bytes)
        if len(binary_string) > 100000:
            n = min(n, 2)
        polar_encoded = binary_string * n
        return polar_encoded, f"Basic Polar encoding (rate 1/{n})"
    
    def encode_data(self, data_bytes, method='direct_binary', **kwargs):
        if method not in self.encoding_methods:
            raise ValueError(f"Unknown encoding method: {method}")
        
        return self.encoding_methods[method](data_bytes, **kwargs)

class FileToBinaryProcessor:
    def __init__(self, input_folder, output_folder):
        self.input_folder = input_folder
        self.output_folder = output_folder
        self.encoder = AdvancedBinaryEncoder()
        self.ensure_directories()
    
    def ensure_directories(self):
        os.makedirs(self.input_folder, exist_ok=True)
        os.makedirs(self.output_folder, exist_ok=True)
    
    def excel_value_to_serializable(self, value):
        """将Excel单元格值转换为可JSON序列化的值"""
        if value is None:
            return ""
        elif isinstance(value, (datetime, date)):
            return value.isoformat()
        elif isinstance(value, (int, float, str, bool)):
            return value
        else:
            try:
                return str(value)
            except:
                return ""
    
    def process_excel_file_optimized(self, file_path, encoding_method='direct_binary'):
        """优化版的Excel处理"""
        print(f"Processing Excel file: {file_path}")
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            all_data = {}
            
            # 限制处理的工作表数量和数据行数
            max_sheets = 3
            max_rows_per_sheet = 1000
            
            sheet_count = 0
            for sheet_name in wb.sheetnames:
                if sheet_count >= max_sheets:
                    print(f"  Warning: Only processing first {max_sheets} sheets")
                    break
                
                sheet = wb[sheet_name]
                sheet_data = []
                row_count = 0
                
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= max_rows_per_sheet:
                        break
                    
                    processed_row = [self.excel_value_to_serializable(cell) for cell in row]
                    if any(cell is not None and cell != "" for cell in processed_row):
                        sheet_data.append(processed_row)
                        row_count += 1
                
                if sheet_data:
                    all_data[sheet_name] = sheet_data
                    sheet_count += 1
            
            # 检查数据量
            estimated_size = len(str(all_data))
            if estimated_size > 10 * 1024 * 1024:
                print(f"  Warning: Data too large ({estimated_size} bytes), truncating...")
                for sheet_name in all_data:
                    all_data[sheet_name] = all_data[sheet_name][:100]
            
            json_str = json.dumps(all_data, ensure_ascii=False, cls=CustomJSONEncoder)
            data_bytes = json_str.encode('utf-8')
            
            print(f"  Data size: {len(data_bytes)} bytes")
            
            binary_string, encoding_info = self.encoder.encode_data(
                data_bytes, encoding_method, chunk_size=500
            )
            
            binary_array = jnp.array([int(bit) for bit in binary_string], dtype=jnp.int32)
            
            return {
                'excel_data': all_data,
                'binary_data': {
                    'binary_string': binary_string,
                    'binary_array': binary_array,
                    'total_bits': len(binary_string),
                    'encoding_method': encoding_method,
                    'encoding_info': encoding_info
                },
                'file_type': 'excel',
                'file_size_bytes': len(data_bytes)
            }
            
        except Exception as e:
            print(f"Error processing Excel file: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def process_word_file(self, file_path, encoding_method='direct_binary'):
        print(f"Processing Word file: {file_path}")
        try:
            doc = Document(file_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_data))
            
            text_content = "\n".join(full_text)
            data_bytes = text_content.encode('utf-8')
            
            binary_string, encoding_info = self.encoder.encode_data(data_bytes, encoding_method)
            binary_array = jnp.array([int(bit) for bit in binary_string], dtype=jnp.int32)
            
            return {
                'original_text': text_content,
                'binary_data': {
                    'binary_string': binary_string,
                    'binary_array': binary_array,
                    'total_bits': len(binary_string),
                    'encoding_method': encoding_method,
                    'encoding_info': encoding_info
                },
                'file_type': 'word',
                'file_size_bytes': len(data_bytes)
            }
            
        except Exception as e:
            print(f"Error processing Word file: {e}")
            return None
    
    def save_binary_data_optimized(self, file_name, processed_data):
        """优化版的二进制数据保存"""
        base_name = os.path.splitext(file_name)[0]
        encoding_method = processed_data['binary_data']['encoding_method']
        
        # 保存二进制数组
        binary_npy_path = os.path.join(self.output_folder, f"{base_name}_{encoding_method}_binary.npy")
        np.save(binary_npy_path, np.array(processed_data['binary_data']['binary_array']))
        
        # 只保存预览
        binary_txt_path = os.path.join(self.output_folder, f"{base_name}_{encoding_method}_preview.txt")
        with open(binary_txt_path, 'w', encoding='utf-8') as f:
            preview_length = min(1000, len(processed_data['binary_data']['binary_string']))
            f.write(processed_data['binary_data']['binary_string'][:preview_length])
            f.write(f"\n... (preview only, total {processed_data['binary_data']['total_bits']} bits)")
            f.write(f"\nFull data saved as: {os.path.basename(binary_npy_path)}")
        
        # 保存编码信息
        json_path = os.path.join(self.output_folder, f"{base_name}_{encoding_method}_info.json")
        json_data = {
            'file_name': file_name,
            'file_type': processed_data['file_type'],
            'file_size_bytes': processed_data['file_size_bytes'],
            'total_bits': processed_data['binary_data']['total_bits'],
            'encoding_method': encoding_method,
            'encoding_info': processed_data['binary_data']['encoding_info'],
            'binary_shape': list(processed_data['binary_data']['binary_array'].shape),
            'binary_dtype': str(processed_data['binary_data']['binary_array'].dtype),
            'note': 'Binary string preview only. Use .npy file for full data.'
        }
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2)
        
        self.generate_statistics_report(base_name, processed_data, encoding_method)
        
        return {
            'binary_preview': binary_txt_path,
            'binary_npy': binary_npy_path,
            'info_json': json_path
        }
    
    def generate_statistics_report(self, base_name, processed_data, encoding_method):
        """生成统计报告"""
        binary_array = np.array(processed_data['binary_data']['binary_array'])
        total_bits = len(binary_array)
        zeros_count = np.sum(binary_array == 0)
        ones_count = np.sum(binary_array == 1)
        zero_ratio = zeros_count / total_bits
        one_ratio = ones_count / total_bits
        
        stats_path = os.path.join(self.output_folder, f"{base_name}_{encoding_method}_statistics.txt")
        with open(stats_path, 'w', encoding='utf-8') as f:
            f.write(f"Binary Data Statistics Report\n")
            f.write(f"File: {base_name}\n")
            f.write(f"Type: {processed_data['file_type']}\n")
            f.write(f"Encoding: {processed_data['binary_data']['encoding_info']}\n")
            f.write(f"Original Size: {processed_data['file_size_bytes']} bytes\n")
            f.write(f"Encoded Bits: {total_bits}\n")
            f.write(f"Coding Rate: {processed_data['file_size_bytes'] * 8 / total_bits:.4f}\n\n")
            f.write(f"Bit Distribution:\n")
            f.write(f"  Zeros: {zeros_count} ({zero_ratio:.4f})\n")
            f.write(f"  Ones: {ones_count} ({one_ratio:.4f})\n")
    
    def process_all_files(self, encoding_method='direct_binary'):
        processed_files = []
        
        print(f"Using encoding method: {encoding_method}")
        
        for file_name in os.listdir(self.input_folder):
            file_path = os.path.join(self.input_folder, file_name)
            
            if os.path.isfile(file_path):
                processed_data = None
                
                if file_name.endswith('.docx'):
                    processed_data = self.process_word_file(file_path, encoding_method)
                elif file_name.endswith('.xlsx'):
                    processed_data = self.process_excel_file_optimized(file_path, encoding_method)
                
                if processed_data:
                    output_files = self.save_binary_data_optimized(file_name, processed_data)
                    processed_files.append({
                        'input_file': file_name,
                        'output_files': output_files,
                        'processing_data': processed_data
                    })
                    print(f"Processed: {file_name}")
        
        return processed_files

def main():
    input_folder = "C:/SJTU-simcom/FYPSignal/input_files"
    output_folder = "C:/SJTU-simcom/FYPSignal/output_files"
    
    processor = FileToBinaryProcessor(input_folder, output_folder)
    
    encoding_methods = list(processor.encoder.encoding_methods.keys())
    
    print("Available encoding methods:")
    for i, method in enumerate(encoding_methods, 1):
        print(f"{i}. {method}")
    
    try:
        choice = input(f"Select encoding method (1-{len(encoding_methods)}, default=1): ").strip()
        choice = int(choice) if choice else 1
        selected_method = encoding_methods[choice-1]
    except:
        selected_method = 'direct_binary'
    
    print(f"Starting file processing with {selected_method} encoding...")
    results = processor.process_all_files(selected_method)
    print(f"Processing completed. Total files processed: {len(results)}")

if __name__ == "__main__":
    main()